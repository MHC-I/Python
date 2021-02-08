from docx import Document
import xlwt
import os

'''
1.不能直接处理doc文件（当脚本报错停止时，注意检查文件是否为docx）
'''
#在内存中创建一个空Excel
# 创建一个workbook 设置编码
workbook = xlwt.Workbook(encoding = 'utf-8')
# 创建一个worksheet
worksheet = workbook.add_sheet('My Worksheet')


#读取word中相应数据，处理后添加到Excel中
def word_to_excel(file_path):
    document=Document(file_path)#读取word
    tables=document.tables #获取文件中的表格集
    table=tables[0 ]#获取文件中的第一个表格
    #处理学号
    number= document.tables[0].rows[0].cells[8]
    number=number.text.replace(' ','')
    #处理姓名
    name= document.tables[0].rows[0].cells[1]
    name=name.text.replace(' ','')
    #处理班级
    grade=document.tables[0].rows[1].cells[3]
    grade=grade.text.replace(' ','')
    #处理日期
    date=document.tables[0].rows[3].cells[4]
    date=date.text.replace(' ','')
    date=date.split('日',1)
    date=date[0]+'日'

    # 写入excel
    # 参数对应 行, 列, 值
    worksheet.write(travel,0, label = str(grade))
    worksheet.write(travel,1, label = str(name))
    worksheet.write(travel,2, label = str(number))
    worksheet.write(travel,3, label = str(date))
    print(grade+name+"已导入")

#Excel起始行
travel=0
#文件路径
path='/Volumes/CoCo/确认表'
#用walk方法遍历路径下的所有文件
for root,dirs,files in os.walk(path):
    files = [f for f in files if not f[0] == '.']#忽略隐藏文件
    dirs[:] = [d for d in dirs if not d[0] == '.']
    for name in files:
        file_path=os.path.join(root,name)
        word_to_excel(file_path)
        travel=travel+1
# 保存Excel文件到脚本目录
workbook.save('docx_to_excel.xls')